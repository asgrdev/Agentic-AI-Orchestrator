"""
Document Processing Skills
Skills for reading, parsing, and analyzing various document formats
Uses docling and other advanced methods
"""

import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import json

logger = logging.getLogger(__name__)


# ============================================================================
# DOCUMENT READING SKILLS
# ============================================================================

def read_pdf_with_docling(file_path: str, extract_tables: bool = True, 
                          extract_images: bool = False) -> Dict[str, Any]:
    """
    Read PDF using docling for advanced extraction
    
    Args:
        file_path: Path to PDF file
        extract_tables: Extract tables from PDF
        extract_images: Extract images from PDF
    
    Returns:
        Dict with text, tables, images, and metadata
    """
    try:
        from ingestion.docling_processor import DoclingProcessor
        
        processor = DoclingProcessor()
        result = processor.process_document(
            file_path,
            extract_tables=extract_tables,
            extract_images=extract_images
        )
        
        return {
            "success": True,
            "file_path": file_path,
            "text": result.get("text", ""),
            "tables": result.get("tables", []),
            "images": result.get("images", []),
            "metadata": result.get("metadata", {}),
            "page_count": result.get("page_count", 0),
            "method": "docling"
        }
        
    except ImportError:
        # Fallback to PyMuPDF
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            
            return {
                "success": True,
                "file_path": file_path,
                "text": text,
                "tables": [],
                "images": [],
                "metadata": {},
                "page_count": len(doc),
                "method": "pymupdf_fallback"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to read PDF: {str(e)}",
                "file_path": file_path
            }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def read_docx_file(file_path: str, extract_tables: bool = True) -> Dict[str, Any]:
    """
    Read DOCX file with table extraction
    
    Args:
        file_path: Path to DOCX file
        extract_tables: Extract tables from document
    
    Returns:
        Dict with text, tables, and metadata
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Extract tables
        tables = []
        if extract_tables:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
        
        return {
            "success": True,
            "file_path": file_path,
            "text": text,
            "tables": tables,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(tables),
            "method": "python-docx"
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "python-docx not installed. Install with: pip install python-docx",
            "file_path": file_path
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def read_excel_file(file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
    """
    Read Excel file with pandas
    
    Args:
        file_path: Path to Excel file
        sheet_name: Specific sheet to read (None for all sheets)
    
    Returns:
        Dict with data, sheets, and metadata
    """
    try:
        import pandas as pd
        
        if sheet_name:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            data = {
                sheet_name: df.to_dict('records')
            }
            sheets = [sheet_name]
        else:
            excel_file = pd.ExcelFile(file_path)
            sheets = excel_file.sheet_names
            data = {}
            for sheet in sheets:
                df = pd.read_excel(file_path, sheet_name=sheet)
                data[sheet] = df.to_dict('records')
        
        return {
            "success": True,
            "file_path": file_path,
            "data": data,
            "sheets": sheets,
            "sheet_count": len(sheets),
            "method": "pandas"
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "pandas not installed. Install with: pip install pandas openpyxl",
            "file_path": file_path
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def read_csv_file(file_path: str, delimiter: str = ",", encoding: str = "utf-8") -> Dict[str, Any]:
    """
    Read CSV file with pandas
    
    Args:
        file_path: Path to CSV file
        delimiter: CSV delimiter
        encoding: File encoding
    
    Returns:
        Dict with data and metadata
    """
    try:
        import pandas as pd
        
        df = pd.read_csv(file_path, delimiter=delimiter, encoding=encoding)
        
        return {
            "success": True,
            "file_path": file_path,
            "data": df.to_dict('records'),
            "columns": list(df.columns),
            "row_count": len(df),
            "column_count": len(df.columns),
            "method": "pandas"
        }
        
    except ImportError:
        # Fallback to built-in csv
        try:
            import csv
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                data = list(reader)
            
            return {
                "success": True,
                "file_path": file_path,
                "data": data,
                "columns": list(data[0].keys()) if data else [],
                "row_count": len(data),
                "method": "csv_fallback"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def read_markdown_file(file_path: str, parse_structure: bool = True) -> Dict[str, Any]:
    """
    Read Markdown file with structure parsing
    
    Args:
        file_path: Path to Markdown file
        parse_structure: Parse headings and structure
    
    Returns:
        Dict with text, structure, and metadata
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        structure = []
        if parse_structure:
            lines = content.split('\n')
            for i, line in enumerate(lines, 1):
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    structure.append({
                        "line": i,
                        "level": level,
                        "title": title
                    })
        
        return {
            "success": True,
            "file_path": file_path,
            "text": content,
            "structure": structure,
            "line_count": len(content.split('\n')),
            "heading_count": len(structure),
            "method": "built-in"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def detect_file_type(file_path: str) -> Dict[str, Any]:
    """
    Detect file type and suggest appropriate reader
    
    Args:
        file_path: Path to file
    
    Returns:
        Dict with file type, mime type, and suggested reader
    """
    try:
        import mimetypes
        
        path = Path(file_path)
        extension = path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Map extensions to readers
        reader_map = {
            '.pdf': 'read_pdf_with_docling',
            '.docx': 'read_docx_file',
            '.doc': 'read_docx_file',
            '.xlsx': 'read_excel_file',
            '.xls': 'read_excel_file',
            '.csv': 'read_csv_file',
            '.md': 'read_markdown_file',
            '.txt': 'read_text_file',
            '.json': 'read_json_file',
            '.xml': 'read_xml_file',
            '.html': 'read_html_file',
        }
        
        suggested_reader = reader_map.get(extension, 'read_text_file')
        
        # Get file stats
        stats = path.stat()
        
        return {
            "success": True,
            "file_path": file_path,
            "extension": extension,
            "mime_type": mime_type,
            "suggested_reader": suggested_reader,
            "file_size": stats.st_size,
            "file_size_mb": stats.st_size / (1024 * 1024),
            "exists": path.exists(),
            "is_file": path.is_file()
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


def extract_text_from_image(image_path: str, language: str = "eng") -> Dict[str, Any]:
    """
    Extract text from image using OCR
    
    Args:
        image_path: Path to image file
        language: OCR language (default: English)
    
    Returns:
        Dict with extracted text and confidence
    """
    try:
        from PIL import Image
        import pytesseract
        
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang=language)
        
        # Get confidence data
        data = pytesseract.image_to_data(image, lang=language, output_type=pytesseract.Output.DICT)
        confidences = [int(conf) for conf in data['conf'] if conf != '-1']
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        return {
            "success": True,
            "image_path": image_path,
            "text": text,
            "confidence": avg_confidence,
            "word_count": len(text.split()),
            "method": "tesseract"
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "pytesseract not installed. Install with: pip install pytesseract pillow",
            "image_path": image_path
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "image_path": image_path
        }


def batch_process_documents(file_paths: List[str], auto_detect: bool = True) -> Dict[str, Any]:
    """
    Process multiple documents in batch
    
    Args:
        file_paths: List of file paths
        auto_detect: Auto-detect file types
    
    Returns:
        Dict with results for each file
    """
    results = []
    
    for file_path in file_paths:
        try:
            if auto_detect:
                detection = detect_file_type(file_path)
                if not detection['success']:
                    results.append({
                        "file_path": file_path,
                        "success": False,
                        "error": "Failed to detect file type"
                    })
                    continue
                
                reader_name = detection['suggested_reader']
                reader_func = globals().get(reader_name)
                
                if reader_func:
                    result = reader_func(file_path)
                    results.append(result)
                else:
                    results.append({
                        "file_path": file_path,
                        "success": False,
                        "error": f"No reader found for {detection['extension']}"
                    })
            else:
                # Try to read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                results.append({
                    "file_path": file_path,
                    "success": True,
                    "text": content,
                    "method": "text_fallback"
                })
                
        except Exception as e:
            results.append({
                "file_path": file_path,
                "success": False,
                "error": str(e)
            })
    
    successful = sum(1 for r in results if r.get('success', False))
    
    return {
        "success": True,
        "total_files": len(file_paths),
        "successful": successful,
        "failed": len(file_paths) - successful,
        "results": results
    }


# Register skills
DOCUMENT_SKILLS = {
    "read_pdf_with_docling": {
        "function": read_pdf_with_docling,
        "category": "DOCUMENT",
        "description": "Read PDF with advanced extraction using docling",
        "parameters": ["file_path", "extract_tables", "extract_images"]
    },
    "read_docx_file": {
        "function": read_docx_file,
        "category": "DOCUMENT",
        "description": "Read DOCX file with table extraction",
        "parameters": ["file_path", "extract_tables"]
    },
    "read_excel_file": {
        "function": read_excel_file,
        "category": "DOCUMENT",
        "description": "Read Excel file with pandas",
        "parameters": ["file_path", "sheet_name"]
    },
    "read_csv_file": {
        "function": read_csv_file,
        "category": "DOCUMENT",
        "description": "Read CSV file with pandas",
        "parameters": ["file_path", "delimiter", "encoding"]
    },
    "read_markdown_file": {
        "function": read_markdown_file,
        "category": "DOCUMENT",
        "description": "Read Markdown with structure parsing",
        "parameters": ["file_path", "parse_structure"]
    },
    "detect_file_type": {
        "function": detect_file_type,
        "category": "DOCUMENT",
        "description": "Detect file type and suggest reader",
        "parameters": ["file_path"]
    },
    "extract_text_from_image": {
        "function": extract_text_from_image,
        "category": "DOCUMENT",
        "description": "Extract text from image using OCR",
        "parameters": ["image_path", "language"]
    },
    "batch_process_documents": {
        "function": batch_process_documents,
        "category": "DOCUMENT",
        "description": "Process multiple documents in batch",
        "parameters": ["file_paths", "auto_detect"]
    }
}


# Auto-register skills
def register_document_skills():
    """Register all document processing skills"""
    from agents.skill_executor import get_global_registry
    
    registry = get_global_registry()
    
    for skill_name, skill_info in DOCUMENT_SKILLS.items():
        registry.register(
            name=skill_name,
            function=skill_info["function"],
            category=skill_info["category"],
            description=skill_info["description"]
        )
    
    logger.info(f"✅ Document processing skills registered: {len(DOCUMENT_SKILLS)} skills")


# Auto-register on import
try:
    register_document_skills()
except Exception as e:
    logger.warning(f"Failed to auto-register document skills: {e}")

